import requests
import json
from io import BytesIO
from docx import Document
import base64
from docx.shared import Inches

def send_post_request(url, data):
    try:
        response = requests.post(url, json=data)
        response.raise_for_status()  # Raise an HTTPError for bad responses
        return response.json()  # Convert response to JSON and return it
    except requests.exceptions.RequestException as e:
        print(f"An error occurred: {e}")
        return None
def download_image(url):
    """
    Downloads an image from a given URL and returns it as a BytesIO object.
    
    :param url: str, URL of the image to download
    :return: BytesIO object of the downloaded image
    """
    try:
        if url.startswith("data:image/"):
            # This is a base64 image
            header, encoded = url.split(",", 1)
            image_data = base64.b64decode(encoded)
            return BytesIO(image_data)
        else:
            # This is a normal URL
            response = requests.get(url)
            if response.status_code == 200:
                return BytesIO(response.content)
            else:
                response.raise_for_status()
    except Exception as e:
        print(f"An error occurred while downloading image: {e}")
        return None
def get_Option(question):
    try:
        A = B = C = D = Aexp = Bexp = Cexp = Dexp = isCorrect = None
        for Option in question["Options"]:
            if(Option["OptionLetter"]=="A"):
                A=Option["OptionText"]
                if "ExplanationText" in Option:
                    Aexp=Option["ExplanationText"]
            elif(Option["OptionLetter"]=="B"):
                B=Option["OptionText"]
                if "ExplanationText" in Option:
                    Bexp=Option["ExplanationText"]
            elif(Option["OptionLetter"]=="C"):
                C=Option["OptionText"]
                if "ExplanationText" in Option:
                    Cexp=Option["ExplanationText"]
            elif(Option["OptionLetter"]=="D"):
                D=Option["OptionText"]
                if "ExplanationText" in Option:
                    Dexp=Option["ExplanationText"]
            if(Option["IsCorrect"]==True):
                isCorrect=Option["OptionLetter"]
    except Exception as e:
            print(f"An error occurred: {e}")


    return [A,B,C,D,Aexp,Bexp,Cexp,Dexp,isCorrect]
def get_question_data(question):
    try:
        QuestionText=QuestionImage=ExplanationText=ExplantionImage=None
        QuestionText=question["QuestionText"]
        if(question["ExplanationText"]!=""):
            ExplanationText=question["ExplanationText"]
        if(question["QuestionImage"]!=""):
            QuestionImage=download_image(question["QuestionImage"])
        if(question["ExplanationImage"]!=""):
            ExplantionImage=question["ExplanationImage"]
    except Exception as e:
        print("Exception Occur during Question Data",e)
    return [QuestionText,QuestionImage,ExplanationText,ExplantionImage]
def create_word_document(document,question_data,option_data):
    try:
        document.add_paragraph(f"Q:){question_data[0]}")
        if question_data[1]:
            document.add_picture(question_data[1], width=Inches(4))
        document.add_paragraph(f":Type: S")
        document.add_paragraph(f"A:) {option_data[0]}")
        document.add_paragraph(f"B:) {option_data[1]}")
        document.add_paragraph(f"C:) {option_data[2]}")
        document.add_paragraph(f"D:) {option_data[3]}")
        document.add_paragraph(f":Correct: {option_data[8]}")
        document.add_paragraph(f":MsgCorrect: {option_data[4]}")
        document.add_paragraph(f":MsgIncorrect: {question_data[2]}")
        if question_data[3]:
            document.add_picture(question_data[3], width=Inches(4))
    except Exception as e:
        print("Error Occur whike creating  Word Docs")

    

    
        # Add some space between questions
    document.add_paragraph()


topics = [
    'Atomic Structure Practice',
]

results = []
index=0
for topic in topics:
    document=Document()
    path=f"{index}-{topic}.docx"
    index=index+1
    for i in range(1, 50):
        url = f"https://prodapi.premed.pk/api/decks/getlatestdeckquestion/{i}"
        print("Round No: ",i)

        data_to_send = {"DeckName": topic}
        result = send_post_request(url, data_to_send)
        if result:
            if len(result["questions"])>0:
                #print(f"Received response: {result}")
                index2=0
                for question in result["questions"]:
                    index2=index2+1
                    print("Question No: ",index2)
                    questiondata=get_question_data(question)
                    OptionData=get_Option(question)
                    create_word_document(document,questiondata,OptionData)

                    
                # print(i)
                # results.append(result)
            else:
                break
        else:
            print(f"Failed to get a valid response from {url}")
            break
    document.save(path)
#print(f"All results: {results}")

# # Writing the results to a file
# with open("result.txt", 'w') as file:
#     file.write(json.dumps(results, indent=4))
