import requests
import re
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO
import base64


ProgramId="78cae060-da0f-4ea3-9512-48722a9b2c10"
UserId="c1294025-6bb4-4c2e-9b46-0511ed02c6c3"

def makePostRequest(url,data):
    response = requests.post(url=url, json=data)

# Check if the request was successful
    if response.status_code == 200:
        # Parse the JSON response
        json_response = response.json()
        return json_response
    else:
        print(f"Failed to send POST request.on Url:  {url} Status code: {response.status_code}")
        return None
def cut_after_ism(url):
    if ".ism" in url:
        return url.split(".ism")[0] + ".ism"
    else:
        return url
def create_word_document(document,questionText,correctOption):
    try:
        document.add_paragraph(f"Q:)")
        document.add_picture(questionText, width=Inches(4))
        document.add_paragraph(f":Type: S")
        document.add_paragraph(f"A:)A")
        document.add_paragraph(f"B:)B")
        document.add_paragraph(f"C:)C")
        document.add_paragraph(f"D:)D")
        document.add_paragraph(f":Correct: {correctOption}")
        document.add_paragraph(f":MsgCorrect:")
        document.add_paragraph(f":MsgIncorrect:")
    except Exception as e:
        print("Error Occur whike creating  Word Docs")

    

    
        # Add some space between questions
    document.add_paragraph()

def download_image(url):
    try:   # This is a normal URL
        response = requests.get(url)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            response.raise_for_status()
    except Exception as e:
        print(f"An error occurred while downloading image: {e}")
        return None
def create_test(MCQSTest,TestTitle):
    try:
        for subjectTest in MCQSTest:
            MCQS= json.loads(subjectTest['mcqs'])
            CourseId= subjectTest['courseId']
            path=f"{CourseId}-{TestTitle}.docx"
            document=Document()
            for OneMCQs in MCQS:
                CorrectAnswer=OneMCQs['answers'][0]["correctAnswer"]
                print(OneMCQs['questionText'])
                match = re.search(r'src="([^"]+)"',OneMCQs['questionText'] )

        # Extract and print the src attribute if found
                if match:
                    src = match.group(1)
                    print(src)
                else:
                    print("src attribute not found.")
                
                image=download_image(src)
                create_word_document(document,image,CorrectAnswer)
                
            #print(MCQS)
            document.save(path)
    except Exception as e:
        print(e)

        
workbatch=makePostRequest("https://onlinestep.pgc.edu/api/Work/WorkBatch",{"ProgramId": ProgramId, "Batch": 1})
#print(workbatch)
video_linkdata=[]
for days in workbatch:
    try:
        Workdayid=days["workDayId"]
        print(days["fullName"])
        #print(Workdayid)
        Workdaydata=makePostRequest("https://onlinestep.pgc.edu/api/Work/WorkOneExy",{"UserId":UserId,"WorkDayId":Workdayid})
        print(Workdaydata)
        if  Workdaydata[0].get('workSheet'):
            worksheets= json.loads(Workdaydata[0]['workSheet'])
            for worksheet in worksheets:
                print(worksheet["Title"])
        if  Workdaydata[0].get('test'):
            test= json.loads(Workdaydata[0]['test'])
            for tst in test:
                TestTitle=tst["Title"]
                print(TestTitle)
                MCQSTest=makePostRequest("https://onlinestep.pgc.edu/api/Work/GetTest",{"TestId": TestTitle, "WorkDayId": Workdayid})
                #print(MCQSTest)
                create_test(MCQSTest,TestTitle)

        if  Workdaydata[0].get('video'):
            videos = json.loads(Workdaydata[0]["video"])
            for video in videos:
                title=video['Title']
                courseid=video['CourseId']
                videoTitle=f"{courseid}-{title}"
                print(videoTitle)
                url=f"https:{video['VideoLink']}"
                updatedUrl=cut_after_ism(url)
                print(updatedUrl)
                video_linkdata.append({"url":video['VideoLink'],"Title":videoTitle})
            
    except Exception as e:
        print(e)
    

 #print(video_linkdata)   