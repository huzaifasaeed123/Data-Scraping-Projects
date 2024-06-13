import string
import os
os.environ['path'] += r';C:\Program Files\UniConvertor-2.0rc5\dlls'
import cairosvg
import secrets
import requests
from docx import Document
from urllib.parse import urlencode
import base64
from wand.image import Image
from wand.display import display
import tempfile

def login(api_url, username, password):
    headers = {"Content-Type": "application/x-www-form-urlencoded", "Api-Security-Key": "A956FFC5-8069-47F3-B1D7-AC6E11A21BFD"}
    data = {"username": username, "password": password, "grant_type":"password"}
    response = requests.post(api_url, headers=headers, data=urlencode(data))

    if response.status_code == 200:
        login_data = response.json()
        auth_token = login_data.get("access_token")
        return auth_token
    else:
        raise Exception(f"Login failed. Status code: {response.status_code}")

def fetch_subjects(api_url, auth_token):
    headers = {"Authorization": f"Bearer {auth_token}", "Api-Security-Key": "A956FFC5-8069-47F3-B1D7-AC6E11A21BFD"}
    response = requests.get(api_url, headers=headers)

    if response.status_code == 200:
        subjects_data = response.json()
        return subjects_data
    else:
        raise Exception(f"Failed to fetch subjects data from API. Status code: {response.status_code}")

def fetch_questions(api_url, auth_token):
    headers = {"Authorization": f"Bearer {auth_token}", "Api-Security-Key": "A956FFC5-8069-47F3-B1D7-AC6E11A21BFD"}
    response = requests.get(api_url, headers=headers)

    if response.status_code == 200:
        questions_data = response.json()
        return questions_data
    else:
        raise Exception(f"Failed to fetch questions data from API. Status code: {response.status_code}")
    
def download_image(image_url):
    response = requests.get(image_url)

    if response.status_code == 200:
        return response.content
    else:
        raise Exception(f"Failed to download image from {image_url}. Status code: {response.status_code}")

def upload_image_to_wordpress(wordpress_url, image_data, image_filename):
    headers = {'Authorization': 'Basic ' + wordpress_token.decode('utf-8'), 'Content-Type':'' , "Content-Disposition": "attachment; filename=" + image_filename} # Replace with the appropriate content type if the images are not JPEG format
    response = requests.post(wordpress_url, headers=headers, data=image_data)

    if response.status_code == 201:
        image_url = response.json()["source_url"]
        print(f"Image uploaded to WordPress: {image_url}")
        return image_url
    else:
        raise Exception(f"Failed to upload image to WordPress. Status code: {response.status_code}")
    
def convert_svg_to_png(svg_data):
    svg_file = tempfile.NamedTemporaryFile(suffix=".svg", delete=False)
    png_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    
    try:
        svg_file.write(svg_data)
        svg_file.close()
        
        cairosvg.svg2png(url=svg_file.name, write_to=png_file.name)

        # response = requests.post("https://svgtopng.com/svg-to-png", files={"file": open(svg_file.name, "rb")})
        # if response.status_code == 200:
        #     png_file.write(response.content)
        #     png_file.close()
        #     return png_file.read()
        # else:
        #     raise Exception("Failed to convert SVG to PNG using the online service.")
        png_file.seek(0)
        png_data = png_file.read()
        return png_data
    finally:
        svg_file.close()
        png_file.close()
        os.remove(svg_file.name)
        os.remove(png_file.name)


def replace_option_image(option, wordpress_url):
    def replace_img_src(match):
        N = 7
        image_url = match.group(1)
        print(image_url)
        if "png" in image_url:
            image_data = download_image(image_url)
            image_filename = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for i in range(N)) +    '.png'  # Replace with a unique filename for each image
            uploaded_image_url = upload_image_to_wordpress(wordpress_url, image_data, image_filename)
            return f"<img src=\"{uploaded_image_url}\">"
        elif "svg" in image_url:
            image_data = download_image(image_url)
            png_data=convert_svg_to_png(image_data)
            image_filename = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for i in range(N)) +    '.png'  # Replace with a unique filename for each image
            uploaded_image_url = upload_image_to_wordpress(wordpress_url, png_data, image_filename)
            return f"<img src=\"{uploaded_image_url}\">"
        
        elif "jpg" in image_url:
            image_data = download_image(image_url)
            image_filename = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for i in range(N)) +    '.jpg'  # Replace with a unique filename for each image
            uploaded_image_url = upload_image_to_wordpress(wordpress_url, image_data, image_filename)
            return f"<img src=\"{uploaded_image_url}\">"
        elif "JPG" in image_url:
            image_data = download_image(image_url)
            image_filename = ''.join(secrets.choice(string.ascii_uppercase + string.digits) for i in range(N)) +    '.JPG'  # Replace with a unique filename for each image
            uploaded_image_url = upload_image_to_wordpress(wordpress_url, image_data, image_filename)
            return f"<img src=\"{uploaded_image_url}\">"

    import re
    option = re.sub(r"<img src=\"(.*?)\"", replace_img_src, option)
    return option


def create_word_document(data, output_file):
    document = Document()

    for question in data["TestQuestions"]:
        # Extract the relevant information from the question
        body = replace_option_image(question["Body"], wordpress_url)
        option_a = replace_option_image(question["OptionA"], wordpress_url)
        option_b = replace_option_image(question["OptionB"], wordpress_url)
        option_c = replace_option_image(question["OptionC"], wordpress_url)
        option_d = replace_option_image(question["OptionD"], wordpress_url)
        option_e = replace_option_image(question["OptionE"], wordpress_url)
        
        option_correct = replace_option_image(question["CorrectAnswers"], wordpress_url)
        explanation = replace_option_image(question['Explanation'], wordpress_url)

        document.add_paragraph(f"Q:) {body}")

        # Add the options to the document
        document.add_paragraph(f":Type: S")
        document.add_paragraph(f"A:) {option_a}")
        document.add_paragraph(f"B:) {option_b}")
        document.add_paragraph(f"C:) {option_c}")
        document.add_paragraph(f"D:) {option_d}")
        document.add_paragraph(f"E:) {option_e}")

        document.add_paragraph(f":Correct: {option_correct}")
        document.add_paragraph(f":MsgCorrect: {explanation}")
        document.add_paragraph(f":MsgIncorrect: {explanation}")

        # Add some space between questions
        document.add_paragraph()

    document.save(output_file)

if __name__ == "__main__":
    login_url = "https://api.kipslms.com/api/v5/oauth/token"  # Replace with the login API endpoint URL
    subjects_url = "https://api.kipslms.com/api/v5/CandidateTestService/GetTests?TestFilter=0&SubjectId=29121&CandidateSOSId=503452&PageNumber=1&PageSize=100"  # Replace with the subjects API endpoint URL
    username = "URASHID231785"  # Replace with your username
    password = "172804"  # Replace with your password

    wordpress_user = "abaid987"
    wordpress_password = "jWKY daN1 KvS6 fsZl sz5A qlgS"

    wordpress_url = "https://saeedmdcatlms.com/wp-json/wp/v2/media"

    wordpress_credentials = wordpress_user + ":" + wordpress_password
    wordpress_token = base64.b64encode(wordpress_credentials.encode())

    try:
        auth_token = login(login_url, username, password)
        subjects_data = fetch_subjects(subjects_url, auth_token)
        
        print(auth_token)
        counter = 1
        for subject in subjects_data['Tests']:
            word_output_file = f"PHY/{subject['TestName'].replace('|', '')}.docx"            # Replace with the desired output Word file name
            if os.path.exists(word_output_file):
                word_output_file = f"Testing/{subject['TestName'].replace('|', '')} {counter}.docx"
                counter += 1
            test_id = subject["TestId"]
            questions_url = f"https://api.kipslms.com/api/v5/CandidateTestService/GetAttemptedTestResultDetail?TestId={test_id}&CandidateSOSId=503452"  # Replace with the questions API endpoint URL
            questions_data = fetch_questions(questions_url, auth_token)
            create_word_document(questions_data, word_output_file)
            print("Word document created successfully.")
    except Exception as e:
        print(f"Error: {e}")
